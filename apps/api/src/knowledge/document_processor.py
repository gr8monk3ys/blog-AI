"""
Document processor for the Knowledge Base system.

This module handles parsing documents of various formats (PDF, DOCX, TXT, MD)
and splitting them into semantically meaningful chunks for embedding and retrieval.

Supports:
- PDF parsing with page number tracking
- DOCX parsing with structure preservation
- Plain text and Markdown parsing
- Multiple chunking strategies (fixed, semantic, paragraph, recursive)
"""

import hashlib
import io
import logging
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from ..types.knowledge import (
    ChunkingConfig,
    ChunkingStrategy,
    ChunkMetadata,
    Document,
    DocumentChunk,
    DocumentMetadata,
    DocumentType,
)

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Exception raised when document processing fails."""

    def __init__(self, message: str, document_name: Optional[str] = None):
        self.document_name = document_name
        super().__init__(message)


class DocumentProcessor:
    """
    Process documents for the knowledge base.

    Handles parsing various document formats and splitting them into
    chunks suitable for embedding and retrieval.
    """

    # Approximate tokens per character (varies by language)
    CHARS_PER_TOKEN = 4

    def __init__(self, chunking_config: Optional[ChunkingConfig] = None):
        """
        Initialize the document processor.

        Args:
            chunking_config: Configuration for chunking behavior.
                           If None, uses default configuration.
        """
        self.config = chunking_config or ChunkingConfig()
        self._validate_config()

    def _validate_config(self) -> None:
        """Validate chunking configuration."""
        if self.config.chunk_size < self.config.min_chunk_size:
            raise ValueError("chunk_size must be >= min_chunk_size")
        if self.config.chunk_size > self.config.max_chunk_size:
            raise ValueError("chunk_size must be <= max_chunk_size")
        if self.config.chunk_overlap >= self.config.chunk_size:
            raise ValueError("chunk_overlap must be < chunk_size")

    def detect_file_type(self, filename: str, content: bytes) -> DocumentType:
        """
        Detect the document type from filename and content.

        Args:
            filename: Name of the file
            content: Raw file content bytes

        Returns:
            Detected document type

        Raises:
            DocumentProcessingError: If file type cannot be determined
        """
        extension = Path(filename).suffix.lower()

        extension_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
        }

        if extension in extension_map:
            return extension_map[extension]

        # Try to detect from content magic bytes
        if content[:4] == b"%PDF":
            return DocumentType.PDF
        if content[:4] == b"PK\x03\x04":  # ZIP header (DOCX)
            return DocumentType.DOCX

        # Default to plain text if content is valid UTF-8
        try:
            content.decode("utf-8")
            return DocumentType.TXT
        except UnicodeDecodeError:
            raise DocumentProcessingError(
                f"Cannot determine file type for: {filename}",
                document_name=filename,
            )

    def parse_document(
        self,
        content: bytes,
        file_type: DocumentType,
        filename: str,
    ) -> Tuple[str, DocumentMetadata]:
        """
        Parse a document and extract text content and metadata.

        Args:
            content: Raw document bytes
            file_type: Type of the document
            filename: Name of the file

        Returns:
            Tuple of (extracted text, document metadata)

        Raises:
            DocumentProcessingError: If parsing fails
        """
        parser_map = {
            DocumentType.PDF: self._parse_pdf,
            DocumentType.DOCX: self._parse_docx,
            DocumentType.TXT: self._parse_text,
            DocumentType.MARKDOWN: self._parse_markdown,
            DocumentType.HTML: self._parse_html,
        }

        parser = parser_map.get(file_type)
        if not parser:
            raise DocumentProcessingError(
                f"Unsupported document type: {file_type}",
                document_name=filename,
            )

        try:
            text, page_count, extra_metadata = parser(content, filename)

            metadata = DocumentMetadata(
                title=extra_metadata.get("title", Path(filename).stem),
                source=filename,
                file_type=file_type,
                file_size_bytes=len(content),
                page_count=page_count,
                author=extra_metadata.get("author"),
                created_date=extra_metadata.get("created_date"),
                modified_date=extra_metadata.get("modified_date"),
                custom_metadata=extra_metadata,
            )

            return text, metadata

        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}", exc_info=True)
            raise DocumentProcessingError(
                f"Failed to parse document: {str(e)}",
                document_name=filename,
            )

    def _parse_pdf(
        self, content: bytes, filename: str
    ) -> Tuple[str, Optional[int], Dict[str, Any]]:
        """Parse a PDF document."""
        try:
            import pypdf
        except ImportError:
            raise DocumentProcessingError(
                "pypdf package not installed. Install with: pip install pypdf",
                document_name=filename,
            )

        try:
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)

            # Extract metadata
            meta = reader.metadata or {}
            extra_metadata = {}
            if meta.title:
                extra_metadata["title"] = meta.title
            if meta.author:
                extra_metadata["author"] = meta.author
            if meta.creation_date:
                extra_metadata["created_date"] = meta.creation_date

            # Extract text with page markers
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Add page marker for citation tracking
                    text_parts.append(f"[PAGE {i + 1}]\n{page_text}")

            full_text = "\n\n".join(text_parts)
            return full_text, len(reader.pages), extra_metadata

        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to parse PDF: {str(e)}",
                document_name=filename,
            )

    def _parse_docx(
        self, content: bytes, filename: str
    ) -> Tuple[str, Optional[int], Dict[str, Any]]:
        """Parse a DOCX document."""
        try:
            import docx
        except ImportError:
            raise DocumentProcessingError(
                "python-docx package not installed. Install with: pip install python-docx",
                document_name=filename,
            )

        try:
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)

            # Extract metadata
            core_props = doc.core_properties
            extra_metadata = {}
            if core_props.title:
                extra_metadata["title"] = core_props.title
            if core_props.author:
                extra_metadata["author"] = core_props.author
            if core_props.created:
                extra_metadata["created_date"] = core_props.created
            if core_props.modified:
                extra_metadata["modified_date"] = core_props.modified

            # Extract text from paragraphs
            text_parts = []
            current_heading = None

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Track headings for section metadata
                if para.style and para.style.name.startswith("Heading"):
                    current_heading = text
                    text_parts.append(f"\n## {text}\n")
                else:
                    text_parts.append(text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)
            return full_text, None, extra_metadata

        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to parse DOCX: {str(e)}",
                document_name=filename,
            )

    def _parse_text(
        self, content: bytes, filename: str
    ) -> Tuple[str, Optional[int], Dict[str, Any]]:
        """Parse a plain text document."""
        try:
            # Try common encodings
            for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
                try:
                    text = content.decode(encoding)
                    return text, None, {}
                except UnicodeDecodeError:
                    continue

            raise DocumentProcessingError(
                "Could not decode text file with any supported encoding",
                document_name=filename,
            )
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to parse text file: {str(e)}",
                document_name=filename,
            )

    def _parse_markdown(
        self, content: bytes, filename: str
    ) -> Tuple[str, Optional[int], Dict[str, Any]]:
        """Parse a Markdown document."""
        text, _, extra_metadata = self._parse_text(content, filename)

        # Extract title from first heading if present
        title_match = re.search(r"^#\s+(.+)$", text, re.MULTILINE)
        if title_match:
            extra_metadata["title"] = title_match.group(1).strip()

        return text, None, extra_metadata

    def _parse_html(
        self, content: bytes, filename: str
    ) -> Tuple[str, Optional[int], Dict[str, Any]]:
        """Parse an HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise DocumentProcessingError(
                "beautifulsoup4 package not installed. Install with: pip install beautifulsoup4",
                document_name=filename,
            )

        try:
            # Decode content
            text_content, _, _ = self._parse_text(content, filename)

            soup = BeautifulSoup(text_content, "html.parser")

            # Extract metadata
            extra_metadata = {}
            title_tag = soup.find("title")
            if title_tag:
                extra_metadata["title"] = title_tag.get_text().strip()

            # Remove script and style elements
            for element in soup(["script", "style", "nav", "footer", "header"]):
                element.decompose()

            # Get text content
            text = soup.get_text(separator="\n")

            # Clean up whitespace
            lines = [line.strip() for line in text.splitlines()]
            text = "\n".join(line for line in lines if line)

            return text, None, extra_metadata

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to parse HTML: {str(e)}",
                document_name=filename,
            )

    def chunk_document(
        self,
        document_id: str,
        text: str,
        metadata: DocumentMetadata,
    ) -> List[DocumentChunk]:
        """
        Split document text into chunks for embedding.

        Args:
            document_id: ID of the parent document
            text: Full document text
            metadata: Document metadata

        Returns:
            List of document chunks
        """
        strategy_map = {
            ChunkingStrategy.FIXED_SIZE: self._chunk_fixed_size,
            ChunkingStrategy.SEMANTIC: self._chunk_semantic,
            ChunkingStrategy.PARAGRAPH: self._chunk_paragraph,
            ChunkingStrategy.SENTENCE: self._chunk_sentence,
            ChunkingStrategy.RECURSIVE: self._chunk_recursive,
        }

        chunker = strategy_map.get(self.config.strategy, self._chunk_recursive)

        # Pre-process: extract page markers if present
        page_map = self._build_page_map(text)

        # Clean text for chunking
        clean_text = self._clean_text(text)

        # Generate chunks
        raw_chunks = chunker(clean_text)

        # Create DocumentChunk objects with metadata
        chunks = []
        for i, (chunk_text, start_char, end_char) in enumerate(raw_chunks):
            # Find page number for this chunk
            page_number = self._find_page_number(start_char, page_map)

            # Find section title if present
            section_title = self._find_section_title(clean_text, start_char)

            chunk_id = f"{document_id}_chunk_{i}"
            token_count = self._estimate_tokens(chunk_text)

            chunk_metadata = ChunkMetadata(
                document_id=document_id,
                chunk_index=i,
                page_number=page_number,
                section_title=section_title,
                start_char=start_char,
                end_char=end_char,
                token_count=token_count,
            )

            chunk = DocumentChunk(
                id=chunk_id,
                content=chunk_text,
                metadata=chunk_metadata,
            )
            chunks.append(chunk)

        # Calculate overlaps
        for i in range(len(chunks) - 1):
            overlap = self._calculate_overlap(
                chunks[i].content, chunks[i + 1].content
            )
            chunks[i].metadata.overlap_with_next = overlap
            chunks[i + 1].metadata.overlap_with_previous = overlap

        logger.info(
            f"Chunked document {document_id} into {len(chunks)} chunks "
            f"using {self.config.strategy.value} strategy"
        )

        return chunks

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text for chunking."""
        # Remove page markers (we've already extracted them)
        text = re.sub(r"\[PAGE \d+\]\n?", "", text)

        # Normalize whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)

        return text.strip()

    def _build_page_map(self, text: str) -> List[Tuple[int, int]]:
        """Build a map of character positions to page numbers."""
        page_map = []
        for match in re.finditer(r"\[PAGE (\d+)\]", text):
            page_num = int(match.group(1))
            char_pos = match.start()
            page_map.append((char_pos, page_num))
        return page_map

    def _find_page_number(
        self, char_pos: int, page_map: List[Tuple[int, int]]
    ) -> Optional[int]:
        """Find the page number for a given character position."""
        if not page_map:
            return None

        for i, (pos, page) in enumerate(page_map):
            if char_pos < pos:
                return page_map[i - 1][1] if i > 0 else 1
        return page_map[-1][1] if page_map else None

    def _find_section_title(self, text: str, char_pos: int) -> Optional[str]:
        """Find the section title that contains this position."""
        # Look for Markdown-style headers
        header_pattern = r"^##?\s+(.+)$"

        # Find the last header before this position
        relevant_text = text[:char_pos]
        headers = list(re.finditer(header_pattern, relevant_text, re.MULTILINE))

        if headers:
            return headers[-1].group(1).strip()
        return None

    def _estimate_tokens(self, text: str) -> int:
        """Estimate the number of tokens in text."""
        return len(text) // self.CHARS_PER_TOKEN

    def _calculate_overlap(self, chunk1: str, chunk2: str) -> int:
        """Calculate token overlap between two chunks."""
        # Simple approximation: find common suffix/prefix
        min_len = min(len(chunk1), len(chunk2))
        max_overlap = min(min_len, self.config.chunk_overlap * self.CHARS_PER_TOKEN)

        for i in range(max_overlap, 0, -1):
            if chunk1[-i:] == chunk2[:i]:
                return i // self.CHARS_PER_TOKEN
        return 0

    def _chunk_fixed_size(self, text: str) -> List[Tuple[str, int, int]]:
        """Split text into fixed-size chunks with overlap."""
        chunk_chars = self.config.chunk_size * self.CHARS_PER_TOKEN
        overlap_chars = self.config.chunk_overlap * self.CHARS_PER_TOKEN

        chunks = []
        start = 0

        while start < len(text):
            end = min(start + chunk_chars, len(text))

            # Try to end at a sentence boundary
            if end < len(text):
                sentence_end = text.rfind(". ", start, end)
                if sentence_end > start + overlap_chars:
                    end = sentence_end + 1

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append((chunk_text, start, end))

            start = end - overlap_chars

        return chunks

    def _chunk_paragraph(self, text: str) -> List[Tuple[str, int, int]]:
        """Split text by paragraphs, merging small ones."""
        paragraphs = re.split(r"\n\n+", text)
        chunks = []
        current_chunk = []
        current_tokens = 0
        start_char = 0
        chunk_start = 0

        for para in paragraphs:
            para_tokens = self._estimate_tokens(para)

            # If paragraph alone exceeds max, split it
            if para_tokens > self.config.max_chunk_size:
                # Flush current chunk first
                if current_chunk:
                    chunk_text = "\n\n".join(current_chunk)
                    chunks.append((chunk_text, chunk_start, start_char))
                    current_chunk = []
                    current_tokens = 0

                # Split large paragraph
                sub_chunks = self._chunk_fixed_size(para)
                for sub_text, sub_start, sub_end in sub_chunks:
                    chunks.append((sub_text, start_char + sub_start, start_char + sub_end))

                chunk_start = start_char + len(para)

            # If adding this paragraph exceeds chunk size, start new chunk
            elif current_tokens + para_tokens > self.config.chunk_size:
                if current_chunk:
                    chunk_text = "\n\n".join(current_chunk)
                    chunks.append((chunk_text, chunk_start, start_char))
                current_chunk = [para]
                current_tokens = para_tokens
                chunk_start = start_char

            else:
                current_chunk.append(para)
                current_tokens += para_tokens

            start_char += len(para) + 2  # +2 for \n\n

        # Flush remaining
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            chunks.append((chunk_text, chunk_start, start_char))

        return chunks

    def _chunk_sentence(self, text: str) -> List[Tuple[str, int, int]]:
        """Split text by sentences, grouping into chunks."""
        # Simple sentence splitting
        sentences = re.split(r"(?<=[.!?])\s+", text)

        chunks = []
        current_chunk = []
        current_tokens = 0
        start_char = 0
        chunk_start = 0

        for sentence in sentences:
            sent_tokens = self._estimate_tokens(sentence)

            if current_tokens + sent_tokens > self.config.chunk_size:
                if current_chunk:
                    chunk_text = " ".join(current_chunk)
                    chunks.append((chunk_text, chunk_start, start_char))
                current_chunk = [sentence]
                current_tokens = sent_tokens
                chunk_start = start_char
            else:
                current_chunk.append(sentence)
                current_tokens += sent_tokens

            start_char += len(sentence) + 1

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append((chunk_text, chunk_start, start_char))

        return chunks

    def _chunk_semantic(self, text: str) -> List[Tuple[str, int, int]]:
        """
        Split text using semantic boundaries.

        This attempts to preserve meaning by respecting:
        - Paragraph boundaries
        - Section headers
        - List structures
        - Code blocks
        """
        # First, identify semantic boundaries
        boundaries = []

        # Headers (Markdown style)
        for match in re.finditer(r"^#{1,6}\s+.+$", text, re.MULTILINE):
            boundaries.append(("header", match.start()))

        # Paragraph breaks
        for match in re.finditer(r"\n\n+", text):
            boundaries.append(("paragraph", match.start()))

        # List items
        for match in re.finditer(r"^\s*[-*]\s+", text, re.MULTILINE):
            boundaries.append(("list", match.start()))

        # Sort boundaries by position
        boundaries.sort(key=lambda x: x[1])

        # Build chunks respecting boundaries
        chunks = []
        chunk_start = 0
        current_chunk = ""

        for boundary_type, pos in boundaries:
            segment = text[chunk_start:pos]
            segment_tokens = self._estimate_tokens(segment)

            # Always split at headers
            if boundary_type == "header" and current_chunk:
                chunks.append((current_chunk.strip(), chunk_start, pos))
                current_chunk = ""
                chunk_start = pos

            # Check if adding segment exceeds limit
            elif self._estimate_tokens(current_chunk + segment) > self.config.chunk_size:
                if current_chunk:
                    chunks.append((current_chunk.strip(), chunk_start, pos))
                current_chunk = segment
                chunk_start = pos
            else:
                current_chunk += segment

        # Add remaining text
        remaining = text[chunk_start:] if boundaries else text
        if remaining.strip():
            current_chunk += remaining
            if current_chunk.strip():
                chunks.append((current_chunk.strip(), chunk_start, len(text)))

        # Post-process: split any chunks that are still too large
        final_chunks = []
        for chunk_text, start, end in chunks:
            if self._estimate_tokens(chunk_text) > self.config.max_chunk_size:
                sub_chunks = self._chunk_fixed_size(chunk_text)
                for sub_text, sub_start, sub_end in sub_chunks:
                    final_chunks.append((sub_text, start + sub_start, start + sub_end))
            else:
                final_chunks.append((chunk_text, start, end))

        return final_chunks

    def _chunk_recursive(self, text: str) -> List[Tuple[str, int, int]]:
        """
        Recursively split text using decreasing separator granularity.

        This is the default and most robust chunking strategy.
        """

        def split_recursive(
            text: str, separators: List[str], start_offset: int = 0
        ) -> List[Tuple[str, int, int]]:
            """Recursively split text."""
            if not text:
                return []

            text_tokens = self._estimate_tokens(text)

            # If text fits in chunk, return it
            if text_tokens <= self.config.chunk_size:
                return [(text, start_offset, start_offset + len(text))]

            # Find separator to use
            separator = None
            for sep in separators:
                if sep in text:
                    separator = sep
                    break

            if separator is None:
                # No separator found, use fixed-size splitting
                return self._chunk_fixed_size(text)

            # Split by separator
            parts = text.split(separator)
            remaining_seps = separators[separators.index(separator) + 1 :]

            chunks = []
            current_chunk = []
            current_tokens = 0
            current_start = start_offset

            for i, part in enumerate(parts):
                part_tokens = self._estimate_tokens(part)

                # If part alone is too large, recursively split it
                if part_tokens > self.config.chunk_size:
                    # Flush current chunk
                    if current_chunk:
                        chunk_text = separator.join(current_chunk)
                        chunks.append(
                            (chunk_text, current_start, current_start + len(chunk_text))
                        )
                        current_chunk = []
                        current_tokens = 0

                    # Recursively split large part
                    part_start = text.find(part, current_start - start_offset)
                    sub_chunks = split_recursive(
                        part, remaining_seps, start_offset + part_start
                    )
                    chunks.extend(sub_chunks)
                    current_start = start_offset + part_start + len(part)

                # If adding part exceeds limit, start new chunk
                elif current_tokens + part_tokens > self.config.chunk_size:
                    if current_chunk:
                        chunk_text = separator.join(current_chunk)
                        chunks.append(
                            (chunk_text, current_start, current_start + len(chunk_text))
                        )
                    current_chunk = [part]
                    current_tokens = part_tokens
                    part_start = text.find(part, current_start - start_offset)
                    current_start = start_offset + part_start

                else:
                    current_chunk.append(part)
                    current_tokens += part_tokens

            # Flush remaining
            if current_chunk:
                chunk_text = separator.join(current_chunk)
                chunks.append(
                    (chunk_text, current_start, current_start + len(chunk_text))
                )

            return chunks

        return split_recursive(text, self.config.separators)

    def process_document(
        self,
        content: bytes,
        filename: str,
        user_id: str,
        custom_metadata: Optional[Dict[str, Any]] = None,
    ) -> Tuple[Document, List[DocumentChunk]]:
        """
        Process a document end-to-end: parse, chunk, and prepare for embedding.

        Args:
            content: Raw document bytes
            filename: Name of the file
            user_id: ID of the user uploading the document
            custom_metadata: Additional metadata to attach

        Returns:
            Tuple of (Document object, list of DocumentChunks)
        """
        # Detect file type
        file_type = self.detect_file_type(filename, content)

        # Parse document
        text, metadata = self.parse_document(content, file_type, filename)

        # Add custom metadata
        if custom_metadata:
            metadata.custom_metadata.update(custom_metadata)

        # Generate document ID
        doc_id = str(uuid.uuid4())

        # Create document object
        document = Document(
            id=doc_id,
            user_id=user_id,
            filename=filename,
            content=text,
            metadata=metadata,
            status="processing",
        )

        # Chunk document
        chunks = self.chunk_document(doc_id, text, metadata)

        # Update document with chunk count
        document.chunk_count = len(chunks)
        document.status = "ready"

        return document, chunks
