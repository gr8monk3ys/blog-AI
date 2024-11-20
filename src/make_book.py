import os
import json
import time
import argparse
from typing import List, Dict, Optional
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Pt, Length, Inches
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain

# Load environment variables
load_dotenv()

# Check for OpenAI API key
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("OpenAI API key not found. Please set it in your .env file.")

# Pydantic Models
class Topic(BaseModel):
    title: str
    content: Optional[str] = None

class Chapter(BaseModel):
    number: int
    title: str
    topics: List[Topic]

class Book(BaseModel):
    title: str
    chapters: List[Chapter]
    output_file: str = Field(default="book.docx")

class DocumentConfig:
    def __init__(self):
        self.doc = Document()
        self.directory = "content/books/"
        if not os.path.exists(self.directory):
            os.makedirs(self.directory)
        
        # Set document style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(12)
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing = 1.15  # Changed from 1.5 * Pt(12)
        
        # Set page margins
        sections = self.doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

def create_title(topic: str, doc_config: DocumentConfig) -> str:
    try:
        llm = ChatOpenAI(model="gpt-4", temperature=0.9)
        title_template = PromptTemplate(
            input_variables=["topic"],
            template="""Create a compelling and creative book title about {topic}.
            The title should be catchy and memorable, but not longer than 8 words.
            Return only the title, without quotes or any additional text."""
        )
        
        title_chain = LLMChain(
            llm=llm,
            prompt=title_template,
            verbose=True,
        )
        
        title = title_chain.run(topic=topic).strip().strip('"').strip("'")
        
        # Document formatting
        paragraph = doc_config.doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = paragraph.add_run(title)
        title_run.font.size = Pt(24)  # Reduced from 48
        doc_config.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        return title
    except Exception as e:
        print(f"Error generating title: {str(e)}")
        raise

def create_book_structure(title: str) -> Book:
    try:
        llm = ChatOpenAI(model="gpt-4", temperature=0.9)
        structure_template = PromptTemplate(
            input_variables=["title"],
            template="""Create a detailed book structure for a book titled '{title}'.
            Return the structure as a JSON object with exactly 11 chapters, each having exactly 4 topics.
            Use this format:
            {{
                "chapters": [
                    {{
                        "number": 1,
                        "title": "chapter title",
                        "topics": [
                            "topic 1",
                            "topic 2",
                            "topic 3",
                            "topic 4"
                        ]
                    }}
                ]
            }}
            Make sure the chapters and topics flow logically and cover the subject comprehensively."""
        )
        
        structure_chain = LLMChain(
            llm=llm,
            prompt=structure_template,
            verbose=True,
        )
        
        structure_json = json.loads(structure_chain.run(title=title))
        
        # Convert JSON to Pydantic models
        chapters = []
        for chapter_data in structure_json["chapters"]:
            topics = [Topic(title=t) for t in chapter_data["topics"]]
            chapters.append(Chapter(
                number=chapter_data["number"],
                title=chapter_data["title"],
                topics=topics
            ))
        
        return Book(title=title, chapters=chapters)
    except Exception as e:
        print(f"Error generating book structure: {str(e)}")
        raise

def generate_topic_content(topic: Topic, chapter: Chapter, book: Book) -> None:
    try:
        llm = ChatOpenAI(model="gpt-4", temperature=0.7)
        topic_template = PromptTemplate(
            input_variables=["book", "chapter", "topic"],
            template="""Write a detailed section about '{topic}' for chapter '{chapter}' of the book '{book}'.
            The content should be engaging, informative, and approximately 2000 words.
            Focus on providing value to the reader with clear explanations and relevant examples.
            Avoid unnecessary transitions or redundant text.
            Each paragraph should naturally flow into the next."""
        )
        
        topic_memory = ConversationBufferMemory(
            input_key="chapter",
            memory_key="chat_history"
        )
        
        topic_chain = LLMChain(
            llm=llm,
            prompt=topic_template,
            verbose=True,
            output_key="topic",
            memory=topic_memory,
        )
        
        topic.content = topic_chain.run(
            book=book.title,
            chapter=chapter.title,
            topic=topic.title
        ).strip()
        time.sleep(2)
    except Exception as e:
        print(f"Error generating content for topic {topic.title}: {str(e)}")
        raise

def write_book_to_doc(book: Book, doc_config: DocumentConfig) -> None:
    try:
        for chapter in book.chapters:
            # Add chapter title
            doc_config.doc.add_paragraph(f"Chapter {chapter.number}: {chapter.title}\n")
            
            # Add each topic's content
            for topic in chapter.topics:
                if topic.content:
                    doc_config.doc.add_paragraph(topic.content)
            
            # Add page break after each chapter
            doc_config.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    except Exception as e:
        print(f"Error writing book to document: {str(e)}")
        raise

def main():
    try:
        parser = argparse.ArgumentParser(description='Generate a book about a specific topic')
        parser.add_argument('topic', type=str, help='The topic to write about')
        parser.add_argument('--output', type=str, default='book.docx', help='Output file name')
        parser.add_argument('--model', type=str, default='gpt-4', help='OpenAI model to use (default: gpt-4)')
        args = parser.parse_args()
        
        # Initialize document
        doc_config = DocumentConfig()
        
        print("Starting book generation process...")
        print(f"Topic: {args.topic}")
        print(f"Output file: {args.output}")
        print(f"Using model: {args.model}")
        
        # Create book title
        print("\nGenerating book title...")
        title = create_title(args.topic, doc_config)
        print(f"Generated title: {title}")
        
        # Create book structure
        print("\nCreating book structure...")
        book = create_book_structure(title)
        book.output_file = args.output
        
        # Generate content for each topic
        print("\nGenerating content for each chapter...")
        for chapter in book.chapters:
            print(f"\nWorking on Chapter {chapter.number}: {chapter.title}")
            for topic in chapter.topics:
                print(f"  Generating content for topic: {topic.title}")
                generate_topic_content(topic, chapter, book)
        
        # Write to document
        print("\nWriting content to document...")
        write_book_to_doc(book, doc_config)
        
        # Save the document
        output_path = os.path.join(doc_config.directory, book.output_file)
        doc_config.doc.save(output_path)
        print(f"\nBook generated successfully at {output_path}")
        
    except Exception as e:
        print(f"\nError during book generation: {str(e)}")
        raise

if __name__ == "__main__":
    main()