"""DOCX formatter for books."""

import logging
from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from ...exceptions import FormattingError
from ...models import Book
from .base import Formatter

logger = logging.getLogger(__name__)


class DOCXFormatter(Formatter[Book]):
    """
    Formats books as Microsoft Word DOCX documents.

    Includes proper styling, formatting, and pagination.
    """

    def __init__(
        self,
        font_name: str = "Arial",
        font_size: int = 11,
        line_spacing: float = 1.15,
        margins: float = 1.0,
    ):
        """
        Initialize DOCX formatter with styling options.

        Args:
            font_name: Font family name
            font_size: Font size in points
            line_spacing: Line spacing multiplier
            margins: Page margins in inches
        """
        self.font_name = font_name
        self.font_size = font_size
        self.line_spacing = line_spacing
        self.margins = margins

    def format(self, content: Book, **kwargs) -> bytes:
        """
        Format book as DOCX document.

        Args:
            content: Book to format
            **kwargs: Additional options

        Returns:
            DOCX file as bytes

        Raises:
            FormattingError: If formatting fails
        """
        try:
            logger.debug(f"Formatting book: {content.title}")

            # Create document
            doc = Document()

            # Apply styling
            self._apply_document_style(doc)

            # Add title page
            self._add_title_page(doc, content)

            # Add chapters
            for chapter in content.chapters:
                self._add_chapter(doc, chapter)

            # Save to bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(
                f"Book formatted: {len(content.chapters)} chapters, ~{content.word_count} words"
            )

            return buffer.getvalue()

        except Exception as e:
            raise FormattingError(
                f"Failed to format book as DOCX: {e}",
                details={"title": content.title},
            ) from e

    def _apply_document_style(self, doc: Document) -> None:
        """
        Apply global document styling.

        Args:
            doc: Document to style
        """
        # Set normal style
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)

        # Set paragraph format
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(12)
        paragraph_format.space_before = Pt(0)
        paragraph_format.line_spacing = self.line_spacing

        # Set page margins
        for section in doc.sections:
            section.left_margin = Inches(self.margins)
            section.right_margin = Inches(self.margins)
            section.top_margin = Inches(self.margins)
            section.bottom_margin = Inches(self.margins)

    def _add_title_page(self, doc: Document, book: Book) -> None:
        """
        Add title page to document.

        Args:
            doc: Document
            book: Book with title
        """
        # Center-aligned title
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = paragraph.add_run(book.title)
        title_run.font.size = Pt(24)
        title_run.bold = True

        # Add subtitle if present
        if book.subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subtitle_run = subtitle_para.add_run(book.subtitle)
            subtitle_run.font.size = Pt(16)

        # Add author if present
        if book.author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            author_run = author_para.add_run(f"by {book.author}")
            author_run.font.size = Pt(14)

        # Page break after title
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_chapter(self, doc: Document, chapter) -> None:
        """
        Add chapter to document.

        Args:
            doc: Document
            chapter: Chapter to add
        """
        # Chapter heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"Chapter {chapter.number}: {chapter.title}")
        heading_run.font.size = Pt(16)
        heading_run.bold = True

        # Add spacing after heading
        heading.paragraph_format.space_after = Pt(18)

        # Add topic content
        for topic in chapter.topics:
            if topic.content:
                # Add content paragraphs
                # Split on double newlines for paragraph breaks
                paragraphs = topic.content.split("\n\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

        # Page break after chapter
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    @property
    def output_extension(self) -> str:
        """Get file extension."""
        return ".docx"

    @property
    def content_type(self) -> str:
        """Get MIME type."""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
