import os
import json
import docx
import time
import openai
import langchain

from docx import Document
from docx.shared import Pt, Length
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.utilities import WikipediaAPIWrapper
from langchain.memory import ConversationBufferMemory
from langchain.chains import LLMChain, SequentialChain

os.environ["OPENAI_API_KEY"] = apikey
openai.api_key = os.getenv("OPENAI_API_KEY")
doc = Document()
directory = "books/"
if not os.path.exists(directory):
    os.makedirs(directory)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
paragraph_format = style.paragraph_format
paragraph_format.space_after = Length(0)
paragraph_format.space_before = Length(0)
paragraph_format.line_spacing = 1.5 * Pt(12)

def create_title(prompt, doc):
    title_template = PromptTemplate(
        input_variables=["topic"], template="write me a book title about {topic}"
    )

    llm = OpenAI(model_name="gpt-3.5-turbo", temperature=0.9)
    title_memory = ConversationBufferMemory(
        input_key="topic", memory_key="chat_history"
    )
    title_chain = LLMChain(
        llm=llm,
        prompt=title_template,
        verbose=True,
        output_key="title",
        memory=title_memory,
    )
    title = title_chain.run(prompt)

    # Document formatting
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = paragraph.add_run(title)
    title_run.font.size = Pt(48)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    print(f"Title: {title}\n")
    return title


def create_index(title):
    print("Initializing creation of index...")
    index_prompt = f"""Write me a book outline on a book called '{title}' with 11 chapters. Each chapter has 4 topics, output as a 
    json code. Please make sure to not say anything else except output the code. Assuming multiple chapters, it should look exactly 
    like the following in terms of structure: 
        {{\"Title\": \"\",\"Chapters\": [{{\"Chapter 1\": \"\",\"Topics\": [\"\", \"\", \"\"]}},{{\"Chapter 2\": 
        \"\",\"Topics\": [\"\", \"\", \"\"]}},{{\"Chapter 3\": \"\",\"Topics\": [\"\", \"\", \"\"]}}]}}"""

    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=index_prompt,
        temperature=0.9,
        max_tokens=1500,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0.6,
        stop=[" Human:", " AI:"],
    )

    response_dict = response.choices[0].to_dict()
    print(json.dumps(response_dict, indent=4))
    return response_dict


def create_chapter(index, doc):
    json_string = index["text"]
    json_string = json_string[2:]

    data = json.loads(json_string)
    chapters = data["Chapters"]
    title = data["Title"]

    for chapter in chapters:
        chapter_num, chapter_name = list(chapter.items())[0]
        topics = chapter["Topics"]
        # print(f'{chapter_num}: {chapter_name}\n')

        # Document formatting
        doc.add_paragraph(chapter_name + "\n")

        for topic in topics:
            llm = OpenAI(
                model="text-davinci-003",
                temperature=0.7,
                max_tokens=3700,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0,
            )
            topic_template = PromptTemplate(
                input_variables=["book", "chapter", "topic"],
                template=""" The following is a  book called '{book}' that has a chapter named '{chapter}', 
                the section name that needs to be focused on in the chapter is called '{topic}' and must be 2000 words. 
                I don't want transition words such as finally, conclusion, or overall. I don't want spaces between
                paragraphs and the beginning of all paragraphs must be indented:""",
            )

            topic_memory = ConversationBufferMemory(
                input_key="chapter", memory_key="chat_history"
            )
            topic_chain = LLMChain(
                llm=llm,
                prompt=topic_template,
                verbose=True,
                output_key="topic",
                memory=topic_memory,
            )
            topic_text = topic_chain.run(book=title, chapter=chapter_name, topic=topic)
            # print(topic_text)
            time.sleep(2)

            # Document formatting
            doc.add_paragraph(topic_text)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


if __name__ == "__main__":
    prompt = input("Enter the topic here: ")

    if prompt:
        title = create_title(prompt, doc)
        index = create_index(title)
        create_chapter(index, doc)

        # section = doc.sections[0]
        # header = section.header
        # paragraph = header.paragraphs[0]
        # paragraph.text = "Page "
        # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # run = paragraph.add_run()
        # field_code = 'PAGE   \\* MERGEFORMAT'
        # run._r.append(docx.oxml.shared.OxmlElement('w:fldSimple'))
        # run._r[-1].set(docx.oxml.ns.nsdecls('w'), field_code)

        doc.save(directory + "book.docx")
